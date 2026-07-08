"""
Ten moduł wyciąga surowy tekst z załączników CV (PDF lub DOCX), żeby
można było go później wysłać do AI w celu oceny.

Nie próbujemy tutaj "rozumieć" struktury CV (to robi za nas AI w kolejnym
kroku) - wyciągamy po prostu cały tekst z pliku.
"""

import io
import logging

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger("cv_bot")


def extract_text_from_pdf(content: bytes) -> str:
    """Wyciąga tekst ze wszystkich stron pliku PDF."""
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(content: bytes) -> str:
    """Wyciąga tekst ze wszystkich akapitów i tabel pliku DOCX."""
    document = Document(io.BytesIO(content))
    text_parts = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    return "\n".join(text_parts).strip()


def extract_cv_text(filename: str, content: bytes) -> str:
    """Rozpoznaje typ pliku po rozszerzeniu i wyciąga z niego tekst.
    Zwraca pusty string, jeśli plik jest uszkodzony lub pusty (bot nie
    powinien się wtedy wywalić, tylko pominąć takie CV z ostrzeżeniem)."""
    lowered = filename.lower()
    try:
        if lowered.endswith(".pdf"):
            return extract_text_from_pdf(content)
        elif lowered.endswith(".docx"):
            return extract_text_from_docx(content)
        else:
            logger.warning("Nieobsługiwany format pliku: %s", filename)
            return ""
    except Exception as exc:
        logger.error("Błąd podczas wyciągania tekstu z %s: %s", filename, exc)
        return ""
